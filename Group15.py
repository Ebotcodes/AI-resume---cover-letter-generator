import os
import pdfplumber
import requests
import re 
from fpdf import FPDF
from datetime import datetime
import threading
import customtkinter as ctk
from tkinter import filedialog, messagebox
from tkinter.scrolledtext import ScrolledText
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

# Set appearanceG
ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")

# API and model config
TOGETHER_API_KEY = "8e0f28a685bf336d8ce153dd76ea065606cb7905edb0be1fbca5f2d84a031cbf"
LLAMA_MODEL = "mistralai/Mixtral-8x7B-Instruct-v0.1"
FONT_PATH = r"C:\Users\ElijahBabs\Downloads\Roboto (1)\static\Roboto_Condensed-Italic.ttf"

from fpdf import FPDF

class ResumePDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)
        self.set_margins(left=20, top=20, right=20)

        # Add Unicode-compatible Roboto font
        self.add_font("Roboto", "", r"C:\Users\ElijahBabs\Downloads\Roboto (1)\static\Roboto-Regular.ttf", uni=True)
        self.add_font("Roboto", "B", r"C:\Users\ElijahBabs\Downloads\Roboto (1)\static\Roboto-Bold.ttf", uni=True)
        self.set_font("Roboto", size=11)

    def header(self):
        pass

    def add_section(self, title, body):
        self.set_font("Roboto", "B", 11)
        self.cell(0, 10, title, ln=True)
        self.set_font("Roboto", "", 10)
        self.multi_cell(0, 5, body)
        self.ln(4)

    def add_paragraph(self, text, spacing=4):
        self.set_font("Roboto", "", 10)
        self.multi_cell(0, spacing, text)
        self.ln(3)


class ResumeTool:
    def __init__(self, resume_path: str, job_description: str, tone: str = "formal", length: str = "detailed"):
        self.resume_path = resume_path
        self.job_description = job_description
        self.tone = tone
        self.length = length
        self.resume_text = self.extract_resume_text()
        
        
    def extract_resume_text(self) -> str:
        if self.resume_path.lower().endswith(".pdf"):
            return self.extract_text_from_pdf()
        elif self.resume_path.lower().endswith(".docx"):
            return self.extract_text_from_docx()
        

    def extract_text_from_pdf(self) -> str:
        text = ""
        with pdfplumber.open(self.resume_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    
    def extract_text_from_docx(self) -> str:
        doc = Document(self.resume_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def generate_with_llama(self, prompt: str) -> str:
        url = "https://api.together.xyz/inference"
        headers = {
            "Authorization": f"Bearer {TOGETHER_API_KEY}",
            "Content-Type": "application/json"
        }
        data = {
            "model": LLAMA_MODEL,
            "prompt": prompt,
            "max_tokens": 500,
            "temperature": 0.0,
            "top_p": 0.6
        }
        response = requests.post(url, headers=headers, json=data)
        if response.status_code == 200:
            return response.json()["choices"][0]["text"].strip()
        else:
            return f"Error: {response.status_code}\n{response.text}"

    def generate_tailored_resume(self):
            tone_instructions = {
                "formal": "Use a professional tone with clear, concise language.",
                "casual": "Use a relaxed and friendly tone with approachable language.",
                "confident": "Use assertive language that highlights strengths and accomplishments.",
            }

            prompt = f"""
        You are a professional resume writer. Write a custom-tailored resume for the job description below, based on the user's current resume.

        Guidelines:
        - Keep the entire resume concise enough to fit on one page.
        - Use a professional tone with clear, concise language.
        - Do NOT copy the job description.
        - Emphasize relevant experience, skills, and achievements.
        - Organize into proper sections: Summary, Skills, Experience, Education.
        - Use bullet points in Experience, highlighting quantifiable accomplishments.
        - Prioritize the most relevant and recent information.
        - Use concise phrasing to keep the resume brief but impactful.

        [USER RESUME]
        {self.resume_text}

        [JOB DESCRIPTION]
        {self.job_description}

        Begin the one-page tailored resume below:
        """
            raw_result = self.generate_with_llama(prompt)
            return raw_result




    def generate_cover_letter(self) -> str:
        tone_instructions = {
        "formal": "Maintain a professional and respectful tone throughout.",
        "casual": "Write in a relaxed, friendly tone, like a conversation with a colleague.",
        "confident": "Adopt an enthusiastic and assertive tone, showcasing strengths boldly."
    }

        prompt = f"""
        You are an expert career assistant. Write a personalized, {self.tone} cover letter based solely on the applicant's resume and the job description provided. The cover letter must be tailored, written in the first person (I/me/my), and should not copy any sentences or bullet points from the job description.

        IMPORTANT:
        - Keep the entire letter concise enough to fit on one page.
        - DO NOT restate or summarize the job description directly.
        - DO NOT use bullets or lists.
        - DO NOT start the letter with “Dear Hiring Manager” unless there's no company name.
        - Highlight the candidate’s experience and suitability based on the resume.
        - You must write a compelling introduction, a strong middle (skills and experiences), and a polite closing.
        - Output ONLY the cover letter content, nothing else.
        Tone guidance: {tone_instructions.get(self.tone, '')}

        --- RESUME START ---
        {self.resume_text}
        --- RESUME END ---

        --- JOB DESCRIPTION START ---
        {self.job_description}
        --- JOB DESCRIPTION END ---

        Do not copy the original resume. Write a fully rewritten resume, tailored to the job description. Output only the new resume content below this line.
        ========================
        """
        return self.generate_with_llama(prompt)

       

    def clean_generated_resume(self, text: str) -> str:
    # Remove any job description-like sections from the beginning
    # It matches blocks that start with job ad language and removes them
        job_description_patterns = [
        r"We are looking for.*?\.?\n",  # Starts with a common JD phrase
        r"The (ideal )?candidate.*?\.?\n",
        r"Responsibilities include.*?\.?\n",
        r"Job Description:.*?\n",
        r"---\s*Job Description\s*---.*",  # Matches --- Job Description blocks
    ]
        cleaned_text = text

        for pattern in job_description_patterns:
            cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.IGNORECASE | re.DOTALL)

        # Find first known resume section
        sections = ["PROFESSIONAL SUMMARY", "SUMMARY", "OBJECTIVE", "EXPERIENCE", "WORK EXPERIENCE"]
        lines = cleaned_text.splitlines()
        start_idx = next((i for i, line in enumerate(lines) 
                        if any(sec in line.upper() for sec in sections)), 0)
        
        cleaned_lines = [line.strip() for line in lines[start_idx:] if line.strip()]
        return '\n'.join(cleaned_lines)


    def prepare_paragraphs(self, text: str) -> list:
        paragraphs = []
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
            if re.match(r'^[A-Z ]{4,}$', line):  # Likely a section heading
                paragraphs.append(f"## {line.title()} ##")
            else:
                paragraphs.append(line)
        return paragraphs


    def save_to_pdf(self, text: str, path: str):
        pdf = ResumePDF()
        pdf.add_page()

        for para in text.strip().split("\n\n"):
            lines = para.strip().split("\n")
            if len(lines) > 1:
                title = lines[0].strip()
                body = "\n".join(lines[1:]).strip()
                pdf.add_section(title, body)
            else:
                pdf.add_paragraph(para.strip(), spacing=6)

        pdf.output(path)


    def save_to_txt(self, text: str, path: str):
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)

    def save_to_docx(self,text: str, path: str):
        doc = Document()

        # Modify Heading 1 style globally
        heading_style = doc.styles['Heading 1']
        font = heading_style.font
        font.name = 'Roboto'
        font.size = Pt(16)
        font.bold = True

        # Set page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        paragraphs = text.split('\n\n')  # Simple paragraph split

        for para in paragraphs:
            para = para.strip()
            if para.startswith("##") and para.endswith("##"):
                heading_text = para.strip("# ").title()
                p = doc.add_paragraph(heading_text, style='Heading 1')
                # Optional: add a run if you want to set font again
                # run = p.runs[0]
                # run.font.name = 'Roboto'
                # run.font.size = Pt(16)
                # run.font.bold = True
                # rFonts = run._element.rPr.rFonts
                # rFonts.set(qn('w:eastAsia'), 'Roboto')

            elif para.startswith("- ") or para.startswith("• "):
                p = doc.add_paragraph(para[2:].strip(), style='List Bullet')
                run = p.runs[0]
                run.font.name = 'Roboto'
                run.font.size = Pt(11)
                rFonts = run._element.rPr.rFonts
                rFonts.set(qn('w:eastAsia'), 'Roboto')

            else:
                p = doc.add_paragraph(para, style='Normal')
                run = p.runs[0]
                run.font.name = 'Roboto'
                run.font.size = Pt(11)
                rFonts = run._element.rPr.rFonts
                rFonts.set(qn('w:eastAsia'), 'Roboto')

        doc.save(path)


    def evaluate_resume_keywords(self):
        resume_words = set(re.findall(r'\b\w+\b', self.resume_text.lower()))
        job_words = set(re.findall(r'\b\w+\b', self.job_description.lower()))

        # Filter out common stop words to improve relevance
        stop_words = {"the", "and", "a", "to", "in", "of", "for", "with", "on", "at", "is", "are", "as", "an", "by", "from", "or", "be"}
        job_keywords = {word for word in job_words if word not in stop_words}

        matched = resume_words & job_keywords
        missing = job_keywords - resume_words
        score = round(len(matched) / max(len(job_keywords), 1) * 100, 2)

        return {
            "score": score,
            "matched": matched,
            "missing": sorted(missing)
        }


class ResumeApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("AI Resume + Cover Letter Generator")
        self.geometry("950x750")
        self.resume_path = ""
        self.save_folder = ""

        self.create_widgets()
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure(2, weight=1)
        self.grid_rowconfigure(5, weight=1)
        self.grid_rowconfigure(7, weight=1)
        self.progress_bar = ctk.CTkProgressBar(self, mode="indeterminate")
        self.progress_bar.grid(row=1, column=0, columnspan=3, padx=10, pady=5)
        self.progress_bar.grid_remove()
        
    def show_loading(self):
        self.progress_bar.grid()
        self.progress_bar.start()

    def hide_loading(self):
        self.progress_bar.stop()
        self.progress_bar.grid_remove()

    def create_widgets(self):
        ctk.CTkLabel(self, text="Upload Resume:").grid(row=0, column=0, padx=10, pady=10, sticky="ew")
        ctk.CTkButton(self, text="Browse...", command=self.browse_resume).grid(row=0, column=1, sticky="ew")
        self.file_label = ctk.CTkLabel(self, text="No file selected", text_color="gray")
        self.file_label.grid(row=0, column=2, sticky="ew")

        ctk.CTkLabel(self, text="Job Description:").grid(row=1, column=0, padx=10, pady=5, sticky="nw")
        self.job_text = ScrolledText(self, width=100, height=1)
        self.job_text.grid(row=2, column=0, columnspan=3, pady=5, padx=10)

        ctk.CTkButton(self, text="Choose Save Folder", command=self.choose_save_folder).grid(row=3, column=0, pady=5)
        ctk.CTkButton(self, text="Generate Resume", command=self.generate_resume).grid(row=3, column=1, pady=5)
        ctk.CTkButton(self, text="Generate Cover Letter", command=self.generate_cover_letter).grid(row=3, column=2, pady=5)

        ctk.CTkLabel(self, text="Resume Preview:").grid(row=4, column=0, sticky="w", padx=10)
        self.resume_preview = ScrolledText(self, width=100, height=20)
        self.resume_preview.grid(row=5, column=0, columnspan=3, padx=10, pady=5)

        ctk.CTkLabel(self, text="Cover Letter Preview:").grid(row=6, column=0, sticky="w", padx=10)
        self.cover_preview = ScrolledText(self, width=100, height=20)
        self.cover_preview.grid(row=7, column=0, columnspan=3, padx=10, pady=5)
        
        ctk.CTkLabel(self, text="Resume Evaluation:").grid(row=12, column=0, sticky="w", padx=10)
        self.eval_box = ScrolledText(self, width=100, height=4)
        self.eval_box.grid(row=13, column=0, columnspan=3, padx=10, pady=5)

        ctk.CTkButton(self, text="Evaluate Resume", command=self.evaluate_resume).grid(row=14, column=1, pady=5)


        self.format_var = ctk.StringVar(value="PDF")
        self.format_menu = ctk.CTkOptionMenu(self, values=["PDF", "TXT", "DOCX"], variable=self.format_var)
        self.format_menu.grid(row=8, column=1, padx=10)

        ctk.CTkButton(self, text="Download Resume", command=self.save_edited_resume).grid(row=8, column=0, pady=10)
        ctk.CTkButton(self, text="Download Cover Letter", command=self.save_edited_cover_letter).grid(row=8, column=2, pady=10)

        # Tone Option
        ctk.CTkLabel(self, text="Select Tone:").grid(row=9, column=0, sticky="e", padx=10)
        self.tone_var = ctk.StringVar(value="formal")
        self.tone_menu = ctk.CTkOptionMenu(self, values=["formal", "casual", "confident"], variable=self.tone_var)
        self.tone_menu.grid(row=9, column=1, sticky="w")

        # Length Option
        ctk.CTkLabel(self, text="Select Length:").grid(row=10, column=0, sticky="e", padx=10)
        self.length_var = ctk.StringVar(value="detailed")
        self.length_menu = ctk.CTkOptionMenu(self, values=["short", "detailed"], variable=self.length_var)
        self.length_menu.grid(row=10, column=1, sticky="w")

        # Appearance mode toggle
        ctk.CTkLabel(self, text="Appearance Mode:").grid(row=11, column=0, sticky="e", padx=10)
        self.appearance_var = ctk.StringVar(value="System")
        self.appearance_menu = ctk.CTkOptionMenu(
        self, values=["Light", "Dark", "System"],
        variable=self.appearance_var,
        command=self.change_appearance_mode
        )
        self.appearance_menu.grid(row=11, column=1, sticky="nw")
        

    def change_appearance_mode(self, new_mode: str):
        ctk.set_appearance_mode(new_mode)
       

    def browse_resume(self):
        path = filedialog.askopenfilename(filetypes=[("Document Files", ".pdf *.docx"), ("PDF Files", ".pdf"), ("Word Documents", "*.docx")])
        if path:
            self.resume_path = path
            self.file_label.configure(text=os.path.basename(path), text_color="yellow")

    def choose_save_folder(self):
        folder = filedialog.askdirectory(title="Select Folder to Save Files")
        if folder:
            self.save_folder = folder

    def generate_resume(self):
        if not self.resume_path or not self.job_text.get("1.0", "end").strip():
            messagebox.showerror("Error", "Upload a resume and enter job description.")
            return

        def run():
            self.show_loading()
            try:
                tool = ResumeTool(
                self.resume_path,
                self.job_text.get("1.0", "end").strip(),
                tone=self.tone_var.get(),
                length=self.length_var.get()
            )
                result = tool.generate_tailored_resume()
                self.resume_preview.delete("1.0", "end")
                self.resume_preview.insert("1.0", result)
                messagebox.showinfo("Success", "Resume generated successfully!")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to generate resume:\n{e}")
                
            finally:
               self.after(0, self.hide_loading)

        threading.Thread(target=run).start()

    def generate_cover_letter(self):
        if not self.resume_path or not self.job_text.get("1.0", "end").strip():
            messagebox.showerror("Error", "Upload a resume and enter job description.")
            return

        def run():
            self.show_loading()
            try:
                tool = ResumeTool(
                self.resume_path,
                self.job_text.get("1.0", "end").strip(),
                tone=self.tone_var.get(),
                length=self.length_var.get()
            )
                result = tool.generate_cover_letter()
                self.cover_preview.delete("1.0", "end")
                self.cover_preview.insert("1.0", result)
                messagebox.showinfo("Success", "Cover letter generated successfully!")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to generate cover letter:\n{e}")
                
            finally:
                self.after(0, self.hide_loading)

        threading.Thread(target=run).start()

    def save_edited_resume(self):
        if not self.save_folder:
            self.choose_save_folder()
        if self.save_folder:
            text = self.resume_preview.get("1.0", "end").strip()
            if text:
                try:
                    tool = ResumeTool(self.resume_path, self.job_text.get("1.0", "end").strip())
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    format_choice = self.format_var.get()
                    path = os.path.join(self.save_folder, f"edited_resume_{timestamp}.{format_choice.lower()}")
                    getattr(tool, f"save_to_{format_choice.lower()}")(text, path)
                    messagebox.showinfo("Success", f"Resume saved successfully:\n{path}")
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to save resume:\n{e}")

    def save_edited_cover_letter(self):
        if not self.save_folder:
            self.choose_save_folder()
        if self.save_folder:
            text = self.cover_preview.get("1.0", "end").strip()
            if text:
                try:
                    tool = ResumeTool(self.resume_path, self.job_text.get("1.0", "end").strip())
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    format_choice = self.format_var.get()
                    path = os.path.join(self.save_folder, f"edited_cover_letter_{timestamp}.{format_choice.lower()}")
                    getattr(tool, f"save_to_{format_choice.lower()}")(text, path)
                    messagebox.showinfo("Success", f"Cover letter saved successfully:\n{path}")
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to save cover letter:\n{e}")

    def evaluate_resume(self):
        if not self.resume_path or not self.job_text.get("1.0", "end").strip():
            messagebox.showerror("Error", "Upload a resume and enter job description.")
            return

        def run():
            self.show_loading()
            try:
                tool = ResumeTool(
                    self.resume_path,
                    self.job_text.get("1.0", "end").strip(),
                    tone=self.tone_var.get(),
                    length=self.length_var.get()
                )
                result = tool.evaluate_resume_keywords()
                summary = f"Keyword Match Score: {result['score']}%\n\n"
                summary += f"Missing Keywords:\n{', '.join(result['missing']) if result['missing'] else 'None! Great Match!'}"
                self.eval_box.delete("1.0", "end")
                self.eval_box.insert("1.0", summary)
                messagebox.showinfo("Evaluation Complete", "Resume evaluation completed.")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to evaluate resume:\n{e}")
            finally:
                self.after(0, self.hide_loading)

        threading.Thread(target=run).start()

                                  
                                  
if __name__ == "__main__":
    app = ResumeApp()
    app.mainloop()